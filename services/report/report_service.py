import os
import re
from docx import Document
from repositories.supabase_client import supabase
from datetime import datetime
from reportlab.platypus import SimpleDocTemplate, Paragraph, PageBreak, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.pdfgen import canvas as rl_canvas
from reportlab.lib.pagesizes import A4
from reportlab.lib.colors import HexColor, white
from reportlab.lib.enums import TA_JUSTIFY
from reportlab.lib.utils import ImageReader

_LOGO_PATH = os.path.join(
    os.path.dirname(os.path.abspath(__file__)),
    "..", "..", "assets", "klaktify-icon.png"
)

TABLE = "reports"


def _md_to_rl(text: str) -> str:
    """
    Konversi subset Markdown ke format yang bisa dibaca ReportLab Paragraph.

    Transformasi:
      **teks**  → <b>teks</b>    (bold)
      *teks*    → <i>teks</i>    (italic, single asterisk)
      1. teks   → teks           (hapus penomoran di awal baris)
      - teks    → teks           (hapus bullet di awal baris)
      # Judul   → teks           (hapus hash heading)

    ReportLab mendukung subset HTML: <b>, <i>, <u>, <br/>.
    Karakter khusus XML (<, >, &) di-escape agar tidak crash parser.
    """
    if not text:
        return text

    # Escape XML characters DULU sebelum inject tag HTML
    text = text.replace("&", "&amp;")
    text = text.replace("<", "&lt;")
    text = text.replace(">", "&gt;")

    # **bold** → <b>bold</b>  (bold dulu sebelum italic agar **..** tidak terpotong)
    text = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', text, flags=re.DOTALL)

    # *italic* → <i>italic</i>
    text = re.sub(r'\*(.+?)\*', r'<i>\1</i>', text, flags=re.DOTALL)

    # Hapus penomoran di awal baris: "1. ", "2. ", dst
    text = re.sub(r'(?m)^\d+\.\s+', '', text)

    # Hapus bullet di awal baris: "- ", "* "
    text = re.sub(r'(?m)^[-*]\s+', '', text)

    # Hapus hash heading: "# ", "## ", dst
    text = re.sub(r'(?m)^#+\s*', '', text)

    return text.strip()

def generate_docx(summary: str, output_path: str):
    doc = Document()
    
    doc.add_heading("SUMMARY", level=1)
    doc.add_paragraph(summary)

    doc.save(output_path)
    return output_path

def insert_metadata(data: dict, file_path: str):
    payload = {
        "tanggal": data["tanggal"],
        "jam": data["jam"],
        "ruangan": data["ruangan"],
        "kode_matkul": data["kode_matkul"],
        "kode_dosen": data["kode_dosen"],
        "kelas": data["kelas"],
        "transcription_done": True,
        "file_path": file_path,
        "created_at": datetime.utcnow().isoformat()
    }

    supabase.table(TABLE).insert(payload).execute()

def insert_summary_record(report_id: int, file_path: str):
    supabase.table("summary").insert({
        "reports_id": report_id,
        "file_path": file_path
    }).execute()

def generate_pdf(summary: str, output_path: str):
    dir_name = os.path.dirname(output_path)
    if dir_name:
        os.makedirs(dir_name, exist_ok=True)

    doc = SimpleDocTemplate(output_path)
    styles = getSampleStyleSheet()

    content = []
    content.append(Paragraph("<b>SUMMARY</b>", styles["Title"]))
    content.append(Paragraph(summary, styles["BodyText"]))

    doc.build(content)

    return output_path

def _fmt_sec(sec: float) -> str:
    """Konversi detik ke format Xm Ys."""
    sec   = int(sec or 0)
    m, s  = divmod(sec, 60)
    return f"{m}m {s}s" if m else f"{s}s"


def generate_combined_pdf(
    summary: str,
    analysis: dict,
    output_path: str,
    nama_matkul: str = "MATA KULIAH",
    pertemuan_ke: int = 0,
    kode_kelas: str = "",
    tanggal: str = "",
    kode_dosen: str = "",
    kode_matkul: str = "",
    ringkasan: str = "",
    catatan: str = "-",
) -> str:
    """
    Generate PDF laporan 2 halaman — template profesional Clactify:
      Hal 1 — Header · Info Card · Section A (Ringkasan) · Section B (Kesesuaian RPS)
      Hal 2 — Section C (Aktivitas + Progress Bar) · Section D (Kesesuaian Metode)

    Menggunakan ReportLab Canvas langsung (bukan Platypus) untuk kontrol layout penuh.
    """
    dir_name = os.path.dirname(output_path)
    if dir_name:
        os.makedirs(dir_name, exist_ok=True)

    PAGE_W, PAGE_H = A4   # 595.28 x 841.89 pt
    ML = 50.0             # margin left
    MR = PAGE_W - 50.0    # margin right
    CW = MR - ML          # content width ≈ 495 pt

    # ── Palette ──────────────────────────────────────────────────────
    C_PRI  = HexColor("#7B1C2E")   # Telkom dark red
    C_GOLD = HexColor("#C8A84B")   # diskusi gold
    C_GRN  = HexColor("#2E7D32")   # sesuai green
    C_LBL  = HexColor("#999999")   # muted label
    C_BG   = HexColor("#F5F5F5")   # card background
    C_BDR  = HexColor("#DDDDDD")   # card border
    C_TXT  = HexColor("#1A1A1A")   # body text
    C_BLU  = HexColor("#1976D2")   # tanya jawab blue

    # ── ParagraphStyle factory ───────────────────────────────────────
    _idx = [0]
    def _ps(**kw):
        _idx[0] += 1
        d = dict(name=f"_clactify_{_idx[0]}", fontName="Helvetica",
                 fontSize=9, leading=13, textColor=C_TXT,
                 spaceAfter=0, spaceBefore=0)
        d.update(kw)
        return ParagraphStyle(**d)

    ST_BODY = _ps(fontSize=9, leading=13, alignment=TA_JUSTIFY)
    ST_CAT  = _ps(fontSize=8.5, leading=13, alignment=TA_JUSTIFY)

    # ── Helper: draw wrapped paragraph, returns new y (bottom) ───────
    def _para(c, txt, st, x, y, w):
        p = Paragraph(txt or "-", st)
        _, h = p.wrap(w, 9999)
        p.drawOn(c, x, y - h)
        return y - h

    # ── Helper: section header bar with letter badge ─────────────────
    def _sec_hdr(c, badge, title, y):
        H = 26
        c.setFillColor(C_PRI)
        c.rect(ML, y - H, CW, H, fill=1, stroke=0)
        bs = 16
        bx, by = ML + 7, y - H + (H - bs) / 2
        c.setFillColor(white)
        c.rect(bx, by, bs, bs, fill=1, stroke=0)
        c.setFillColor(C_PRI)
        c.setFont("Helvetica-Bold", 9)
        c.drawCentredString(bx + bs / 2, by + 4, badge)
        c.setFillColor(white)
        c.setFont("Helvetica-Bold", 10)
        c.drawString(bx + bs + 8, y - H + 9, title)
        return y - H - 6

    # ── Helper: 4-column info card ───────────────────────────────────
    def _info_card(c, matkul, ptm, status, tgl, y):
        H = 48
        cw4 = CW / 4
        c.setFillColor(white)
        c.setStrokeColor(C_BDR)
        c.setLineWidth(0.5)
        c.rect(ML, y - H, CW, H, fill=1, stroke=1)
        hdrs = ["MATA KULIAH", "PERTEMUAN", "STATUS WAKTU", "TANGGAL"]
        vals = [matkul, f"Ke-{ptm}", status, tgl or "-"]
        for i, (h, v) in enumerate(zip(hdrs, vals)):
            cx = ML + i * cw4 + 8
            if i > 0:
                c.setStrokeColor(C_BDR)
                c.line(ML + i * cw4, y - 3, ML + i * cw4, y - H + 3)
            c.setFillColor(C_LBL)
            c.setFont("Helvetica", 7)
            c.drawString(cx, y - 13, h)
            fc = (C_GRN if "tepat" in v.lower() else C_PRI) if h == "STATUS WAKTU" else C_TXT
            c.setFillColor(fc)
            c.setFont("Helvetica-Bold", 11)
            c.drawString(cx, y - 34, v)
        return y - H

    # ── Helper: colored square + bold text (status indicator) ────────
    def _badge_text(c, x, y, color, text, fsize=13):
        SZ = 9
        c.setFillColor(color)
        c.rect(x, y + 2, SZ, SZ, fill=1, stroke=0)
        c.setFillColor(color)
        c.setFont("Helvetica-Bold", fsize)
        c.drawString(x + SZ + 5, y, text)

    # ── Helper: single progress bar row ─────────────────────────────
    def _pbar(c, label, pct, color, y):
        LW, bh = 90, 11
        bx = ML + LW
        bw = CW - LW - 28
        c.setFillColor(C_TXT)
        c.setFont("Helvetica", 9)
        c.drawString(ML, y, label)
        c.setFillColor(HexColor("#EBEBEB"))
        c.rect(bx, y - bh + 2, bw, bh, fill=1, stroke=0)
        if pct > 0:
            c.setFillColor(color)
            c.rect(bx, y - bh + 2, bw * pct / 100, bh, fill=1, stroke=0)
        c.setFillColor(C_TXT)
        c.setFont("Helvetica-Bold", 9)
        c.drawRightString(MR, y, f"{pct}%")
        return y - bh - 5

    # ── Helper: box with left accent border (Catatan / Kesimpulan) ───
    def _note_box(c, prefix, text, y):
        PAD, ACC = 10, 4
        iw  = CW - PAD * 2 - ACC
        txt = (f"<b>{prefix}</b> " if prefix else "") + (text or "-")
        p   = Paragraph(txt, ST_CAT)
        _, th = p.wrap(iw, 9999)
        bh  = th + PAD * 2
        c.setFillColor(HexColor("#FFF9F5"))
        c.setStrokeColor(C_BDR)
        c.setLineWidth(0.5)
        c.rect(ML, y - bh, CW, bh, fill=1, stroke=1)
        c.setFillColor(C_PRI)
        c.rect(ML, y - bh, ACC, bh, fill=1, stroke=0)
        p.drawOn(c, ML + ACC + PAD, y - bh + PAD)
        return y - bh

    # ── Helper: footer ───────────────────────────────────────────────
    def _footer(c, pg, tot):
        c.setFillColor(C_LBL)
        c.setFont("Helvetica", 7.5)
        c.drawString(
            ML, 28,
            "Laporan ini dibuat secara otomatis oleh sistem Clactify -- Telkom University"
        )
        c.drawRightString(MR, 28, f"Hal. {pg} / {tot}")

    # ── Helper: number to Indonesian word ───────────────────────────
    def _ordinal(n):
        W = {1:"Satu", 2:"Dua", 3:"Tiga", 4:"Empat", 5:"Lima",
             6:"Enam", 7:"Tujuh", 8:"Delapan", 9:"Sembilan", 10:"Sepuluh",
             11:"Sebelas", 12:"Dua Belas", 13:"Tiga Belas", 14:"Empat Belas",
             15:"Lima Belas", 16:"Enam Belas"}
        try:
            return W.get(int(n), str(n))
        except Exception:
            return str(n)

    # ── Unpack analysis ──────────────────────────────────────────────
    act        = analysis.get("activity") or {}
    kesesuaian = (analysis.get("kesesuaian")         or "-").strip()
    status_wkt = (analysis.get("status_waktu")        or "-").strip()
    ptm_det    = analysis.get("pertemuan_terdeteksi", pertemuan_ke)
    penjelasan = analysis.get("penjelasan",           "-")
    materi_rps = analysis.get("materi_pembelajaran",  "-")
    peng_rps   = analysis.get("pengalaman_rps",       "-")
    met_dom    = analysis.get("metode_dominan",        "-")
    kes_met    = (analysis.get("kesesuaian_metode")   or "-").strip()
    penj_met   = analysis.get("penjelasan_metode",    "-")

    def _pct(k): return int(act.get(k, 0) or 0)
    def _min(k): return round((act.get(k, 0) or 0) / 60)

    cer_pct = _pct("ceramah_pct");     dis_pct = _pct("diskusi_pct")
    tny_pct = _pct("tanya_jawab_pct"); dim_pct = _pct("diam_pct")
    cer_min = _min("ceramah_sec");     dis_min = _min("diskusi_sec")
    tny_min = _min("tanya_jawab_sec"); dim_min = _min("diam_sec")
    tot_min = cer_min + dis_min + tny_min + dim_min
    dominant = (act.get("dominant") or "CERAMAH").upper()

    dom_c = {
        "CERAMAH":    C_PRI,
        "DISKUSI":    C_GOLD,
        "TANYA JAWAB": C_BLU,
        "DIAM":       HexColor("#757575"),
    }.get(dominant, C_PRI)

    is_sesuai = "sesuai" in kesesuaian.lower() and "tidak" not in kesesuaian.lower()
    is_tepat  = "tepat"  in status_wkt.lower()
    is_met_ok = "sesuai" in kes_met.lower()    and "tidak" not in kes_met.lower()

    # ── Create canvas ────────────────────────────────────────────────
    c = rl_canvas.Canvas(output_path, pagesize=A4)

    # ════════════════════════ PAGE 1 ═════════════════════════════════
    y = PAGE_H - 35

    # Header row
    LOGO_H = 13
    LOGO_W = 13
    c.setFillColor(C_PRI)
    c.setFont("Helvetica-Bold", 9)
    c.drawString(ML, y, "UNIVERSITAS TELKOM")
    hdr_txt   = "CLACTIFY -- CLASS ACTIVITY IDENTIFY"
    hdr_txt_w = c.stringWidth(hdr_txt, "Helvetica", 8)
    if os.path.exists(_LOGO_PATH):
        c.drawImage(
            _LOGO_PATH,
            MR - hdr_txt_w - LOGO_W - 4, y - 4,
            width=LOGO_W, height=LOGO_H,
            preserveAspectRatio=True, mask="auto"
        )
    c.setFillColor(C_LBL)
    c.setFont("Helvetica", 8)
    c.drawRightString(MR, y, hdr_txt)
    y -= 22

    # Title
    c.setFillColor(C_TXT)
    c.setFont("Helvetica-Bold", 24)
    c.drawString(ML, y, "LAPORAN IDENTIFIKASI PEMBELAJARAN")
    y -= 16

    # Subtitle — format: FISIKA 2 (AZK1GAB3) | Pertemuan ke-9 | TK-48-GAB1 | MFC
    nm_str  = f"{nama_matkul} ({kode_matkul})" if kode_matkul else nama_matkul
    kls_str = f"  |  {kode_kelas}" if kode_kelas else ""
    kd_str  = f"  |  {kode_dosen}" if kode_dosen else ""
    c.setFillColor(HexColor("#666666"))
    c.setFont("Helvetica", 8.5)
    c.drawString(
        ML, y,
        f"Analisis Kesesuaian & Aktivitas Kelas  |  {nm_str}"
        f"  |  Pertemuan ke-{pertemuan_ke}{kls_str}{kd_str}"
    )
    y -= 10

    # Separator line
    c.setStrokeColor(C_BDR)
    c.setLineWidth(0.5)
    c.line(ML, y, MR, y)
    y -= 12

    # Info card
    y = _info_card(
        c, nama_matkul, pertemuan_ke,
        status_wkt.title() if status_wkt != "-" else "-",
        tanggal, y
    )
    y -= 14

    # ── Section A: Materi Perkuliahan ────────────────────────────────
    y = _sec_hdr(c, "A", "MATERI PERKULIAHAN", y)
    y -= 14

    # Sub-header: Ringkasan
    c.setFillColor(C_PRI)
    c.setFont("Helvetica-Bold", 10)
    c.drawString(ML, y, "Ringkasan Materi Perkuliahan")
    y -= 12

    # Render ringkasan: pembuka + bullet points
    if isinstance(ringkasan, dict):
        pembuka      = ringkasan.get("pembuka", "")
        bullet_items = ringkasan.get("poin", [])
    elif isinstance(ringkasan, list):
        pembuka      = ""
        bullet_items = ringkasan
    else:
        pembuka      = ""
        bullet_items = []

    if pembuka or bullet_items:
        ST_PEMBUKA = _ps(fontSize=9, leading=13)
        ST_BULLET  = _ps(fontSize=9, leading=14)

        if pembuka:
            y = _para(c, pembuka, ST_PEMBUKA, ML, y, CW)
            y -= 8

        for item in bullet_items[:8]:
            txt = f"•  {item.strip()}"
            y   = _para(c, txt, ST_BULLET, ML + 4, y, CW - 4)
            y  -= 3
    else:
        y = _para(c, "............................................", ST_BODY, ML, y, CW)
    y -= 18

    # Sub-header: Detail
    c.setFillColor(C_PRI)
    c.setFont("Helvetica-Bold", 10)
    c.drawString(ML, y, "Detail Materi Perkuliahan")
    y -= 12

    y = _para(c, _md_to_rl(summary or "-"), ST_BODY, ML, y, CW)

    _footer(c, 1, 2)
    c.showPage()

    # ════════════════════════ PAGE 2 ═════════════════════════════════
    y = PAGE_H - 10

    # Running header — top accent line
    c.setFillColor(C_PRI)
    c.rect(ML, y - 4, CW, 4, fill=1, stroke=0)
    y -= 4

    # Running header — text
    c.setFillColor(C_TXT)
    c.setFont("Helvetica-Bold", 7.5)
    c.drawString(
        ML, y - 14,
        f"LAPORAN IDENTIFIKASI PEMBELAJARAN  |  {nama_matkul.upper()}  --  PERTEMUAN KE-{pertemuan_ke}"
    )
    p2_txt   = "CLACTIFY"
    p2_txt_w = c.stringWidth(p2_txt, "Helvetica-Bold", 7.5)
    if os.path.exists(_LOGO_PATH):
        c.drawImage(
            _LOGO_PATH,
            MR - p2_txt_w - LOGO_W - 4, y - 18,
            width=LOGO_W, height=LOGO_H,
            preserveAspectRatio=True, mask="auto"
        )
    c.setFillColor(C_PRI)
    c.setFont("Helvetica-Bold", 7.5)
    c.drawRightString(MR, y - 14, p2_txt)
    y -= 26

    # ── Section B: Kesesuaian RPS ────────────────────────────────────
    y = _sec_hdr(c, "B", "ANALISIS KESESUAIAN RPS", y)
    y -= 10

    GH  = 55
    HW  = CW / 2
    GT  = y

    for col in range(2):
        for row in range(2):
            cx = ML + col * HW
            cy = GT - row * GH
            c.setFillColor(C_BG)
            c.setStrokeColor(C_BDR)
            c.setLineWidth(0.5)
            c.rect(cx, cy - GH, HW, GH, fill=1, stroke=1)

    # Cell: top-left — Materi RPS
    c.setFillColor(C_LBL); c.setFont("Helvetica", 7)
    c.drawString(ML + 8, GT - 13, "MATERI PEMBELAJARAN RPS")
    _para(c, materi_rps, _ps(fontSize=8.5, leading=12),
          ML + 8, GT - 18, HW - 16)

    # Cell: top-right — Kesesuaian materi
    c.setFillColor(C_LBL); c.setFont("Helvetica", 7)
    c.drawString(ML + HW + 8, GT - 13, "KESESUAIAN MATERI")
    _badge_text(c, ML + HW + 8, GT - 40,
                C_GRN if is_sesuai else C_PRI, kesesuaian.upper())

    # Cell: bottom-left — Pertemuan terdeteksi
    c.setFillColor(C_LBL); c.setFont("Helvetica", 7)
    c.drawString(ML + 8, GT - GH - 13, "PERTEMUAN TERDETEKSI")
    c.setFillColor(C_TXT); c.setFont("Helvetica-Bold", 12)
    c.drawString(ML + 8, GT - GH - 35, f"Ke-{ptm_det} ({_ordinal(ptm_det)})")

    # Cell: bottom-right — Status waktu
    c.setFillColor(C_LBL); c.setFont("Helvetica", 7)
    c.drawString(ML + HW + 8, GT - GH - 13, "STATUS WAKTU")
    _badge_text(c, ML + HW + 8, GT - GH - 35,
                C_GRN if is_tepat else C_PRI, status_wkt.upper())

    y = GT - 2 * GH - 10

    # Catatan Evaluator box
    y = _note_box(c, "Catatan Evaluator:", penjelasan, y)
    y -= 8

    # Catatan Materi Tambahan box
    catatan_val = catatan if catatan and catatan.strip() else "-"
    y = _note_box(c, "Catatan Materi Tambahan:", catatan_val, y)
    y -= 14

    # ── Section C: Aktivitas ─────────────────────────────────────────
    y = _sec_hdr(c, "C", "AKTIVITAS PEMBELAJARAN", y)
    y -= 14

    # Dominant activity subtitle
    dom_label = "AKTIVITAS DOMINAN  "
    dw = c.stringWidth(dom_label, "Helvetica", 9)
    c.setFillColor(C_LBL); c.setFont("Helvetica", 9)
    c.drawString(ML, y, dom_label)
    c.setFillColor(dom_c); c.setFont("Helvetica-Bold", 9)
    c.drawString(ML + dw, y, dominant)
    y -= 18

    # Progress bars
    for lbl, pct, col in [
        ("Ceramah",     cer_pct, C_PRI),
        ("Diskusi",     dis_pct, C_GOLD),
        ("Tanya Jawab", tny_pct, C_BLU),
        ("Diam",        dim_pct, HexColor("#BDBDBD")),
    ]:
        y = _pbar(c, lbl, pct, col, y)
    y -= 8

    # Duration summary table — 4 columns
    TH  = 46
    CW4 = CW / 4
    c.setFillColor(C_BG)
    c.setStrokeColor(C_BDR)
    c.setLineWidth(0.5)
    c.rect(ML, y - TH, CW, TH, fill=1, stroke=1)

    for i, (tl, tv, tc, tb) in enumerate([
        ("TOTAL DURASI",       f"{tot_min} Menit",           C_TXT,  True),
        ("CERAMAH",            f"{cer_min} Menit",           C_PRI,  cer_pct > 0),
        ("DISKUSI",            f"{dis_min} Menit",           C_GOLD, dis_pct > 0),
        ("TANYA JAWAB / DIAM", f"{tny_min + dim_min} Menit", C_TXT,  False),
    ]):
        cx = ML + i * CW4 + 8
        if i > 0:
            c.setStrokeColor(C_BDR)
            c.line(ML + i * CW4, y - 3, ML + i * CW4, y - TH + 3)
        c.setFillColor(C_LBL); c.setFont("Helvetica", 7)
        c.drawString(cx, y - 13, tl)
        c.setFillColor(tc)
        c.setFont("Helvetica-Bold" if tb else "Helvetica", 12 if tb else 10)
        c.drawString(cx, y - 34, tv)

    y -= TH + 14

    # ── Section D: Kesesuaian Metode ────────────────────────────────
    y = _sec_hdr(c, "D", "KESESUAIAN METODE PEMBELAJARAN", y)
    y -= 10

    GH3 = 60
    CW3 = CW / 3
    G3T = y

    for i, (lbl, val, vtype) in enumerate([
        ("METODE RPS (RENCANA)",     peng_rps,       "text"),
        ("AKTIVITAS AKTUAL DOMINAN", met_dom,         "bold"),
        ("KESESUAIAN METODE",        kes_met.upper(), "badge"),
    ]):
        cx = ML + i * CW3
        c.setFillColor(C_BG)
        c.setStrokeColor(C_BDR)
        c.setLineWidth(0.5)
        c.rect(cx, G3T - GH3, CW3, GH3, fill=1, stroke=1)
        c.setFillColor(C_LBL); c.setFont("Helvetica", 7)
        c.drawString(cx + 8, G3T - 13, lbl)
        if vtype in ("text", "bold"):
            st = _ps(
                fontSize=8.5, leading=12,
                fontName="Helvetica-Bold" if vtype == "bold" else "Helvetica"
            )
            _para(c, val, st, cx + 8, G3T - 18, CW3 - 16)
        else:
            _badge_text(c, cx + 8, G3T - 40,
                        C_GRN if is_met_ok else C_PRI, val)

    y = G3T - GH3 - 10

    # Kesimpulan box
    y = _note_box(c, "Kesimpulan:", penj_met, y)

    _footer(c, 2, 2)
    c.save()
    return output_path


def get_existing_summary(report_id: int):
    res = supabase.table("reports") \
        .select("summary_path") \
        .eq("id", report_id) \
        .execute()

    if not res.data:
        return None

    row = res.data[0]

    if not row.get("summary_path"):
        return None

    return {
        "file_path": row["summary_path"]
    }
